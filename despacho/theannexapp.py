from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Mm
import xlrd
from datetime import date
import pandas as pd

def set_table_font_size(table, size_pt):
    """
    Cambia el tamaño de fuente de todo el texto en una tabla.
    :param table: objeto Table de python-docx
    :param size_pt: tamaño en puntos (int o float)
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)

def armar_esqueleto_dos(documento, oficina, encabezados_excel, nombre_expediente = None, separar = False):
    """
    Arma la hoja con determinado formato del word. Y crea la tabla a completar.
    
    :param documento: Documento que se está escribiendo
    :param oficina: [int, str], array con el número de la oficina y el nombre de la oficina
    :param encabezados_excel: list[str] lista con los nombres de los encabezados del excel
    """

    # El nombre del anexo correspondiente a la oficina es del tipo "oficina - nombre de la oficina"
    if nombre_expediente and not separar:
        parrafo_exp = documento.add_paragraph()
        run_exp = parrafo_exp.add_run(nombre_expediente)
        run_exp.bold = True
        run_exp.underline = True
        run_exp.font.size = Pt(16)
        parrafo_exp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    parrafo_exp = documento.add_paragraph()
    nombre_anexo = str(oficina[0]) + " - " + oficina[1]
    run_exp = parrafo_exp.add_run(nombre_anexo)
    run_exp.bold = True
    run_exp.underline = True
    run_exp.font.size = Pt(16)
    parrafo_exp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Tabla
    tabla = documento.add_table(rows=1, cols=len(encabezados_excel) - 1) # El número de columnas tiene que ser la cantidad de columnas del excel, menos las dos
    # que representan el número de oficina y el nombre de la oficina, y una más que es el espacio para firmar
    tabla.style = 'Table Grid'  # Bordes visibles

    encabezado = tabla.rows[0].cells
    idx_col = 0
    for col_excel in encabezados_excel:
        # Iteramos sobre los nombres de las columnas del archivo excel para copiarlas en la tabla
        if idx_col == 0 or idx_col == 1: #Estas dos iteraciones hacen referencia al número y nombre de la oficina
            idx_col += 1
            continue
        else:
            encabezado[idx_col - 2].text = col_excel
        idx_col += 1
       
    encabezado[idx_col - 2].text = "NOTIFICACIÓN, FIRMA Y FECHA" #Espacio en la última columna para la firma
        
    return tabla

def armar_anexo_dosV2(planilla, separar):
    """
    Pasa los datos de la planilla de excel a un formato tabla en un Word.

    :param planilla: Planilla  archivo .xlsx de la cuál se están sacando los datos.
    :param separar: str si es true, devolvemos un archivo por oficina
    """
    nombre_expediente = planilla.name.split("xls")[0]
    documento = crear_documento(nombre_expediente)
    wb = xlrd.open_workbook(file_contents = planilla.read())
    ws = wb.sheet_by_index(0)
    encabezado = ws.row_values(0) #los encabezados de planilla de excel

    oficina_anterior_num = str(int(ws.cell_value(1, 0))) # Primer número de oficina de la planilla excel
    oficina_anterior_nom = str(ws.cell_value(1,1)) # Primer nombre de oficina de la planilla excel
    numero_oficina = str(int(ws.cell_value(1, 0)))
    nombre_oficina = str(ws.cell_value(1, 1))
    tabla = armar_esqueleto_dos(documento, [numero_oficina, nombre_oficina], encabezado, separar = separar)

    documentos = [] #array de documentos creados

    for row_idx in range(1, ws.nrows): #itero sobre las filas

        row = ws.row(row_idx)
        
        numero_oficina = str(int(row[0].value))
        nombre_oficina = str(row[1].value)

        if numero_oficina != oficina_anterior_num and nombre_oficina != oficina_anterior_nom:
            #Si me encuentro con una oficina distinta tengo dos caminos: si se decide separar, 
            #creamos un documento nuevo, sino, seguimos en el nuevo documento pero con una hoja aparte
            
            if separar:
                documentos.append(documento)
                documento = crear_documento(nombre_expediente)
                
            else:
                 documento.add_page_break()

            tabla = armar_esqueleto_dos(documento, [numero_oficina, nombre_oficina], encabezado, nombre_expediente = nombre_expediente, separar = separar)
            oficina_anterior_num = numero_oficina
            oficina_anterior_nom = nombre_oficina
            
        fila = tabla.add_row().cells
    
        for i, cell in enumerate(row):
            
            valor_celda = cell.value
            if i == 0 or i == 1: #una precondición del programa es que las primeras dos columnas refieren a nro de oficina, y oficina. 
                continue
            elif cell.ctype == xlrd.XL_CELL_DATE: #si el tipo de celda es fecha, xlrd lo devuelve como un float, así que hay que convertirlo a fecha
                tupla_fecha = xlrd.xldate.xldate_as_tuple(valor_celda, datemode = 0) #(year, month, day, hour, minute, second)
                date_obj = date(*tupla_fecha[:3])
                # 2. Format as DD/MM/YYYY
                formatted_date = date_obj.strftime("%d/%m/%Y")
                fila[i - 2].text = formatted_date
            elif isinstance(valor_celda,str):
                fila[i - 2].text = str(valor_celda) if valor_celda else ""
            elif isinstance(valor_celda,float):
                fila[i - 2].text = str(int(valor_celda)) if valor_celda else "" #para que nos nos aparezcan decimales

    documentos.append(documento)
    return documentos


def crear_documento(nombre_expediente):
    '''
    Función para crear el documento Word y darle el formato deseado

    :param nombre_expediente: número del expediente para que sea el titulo de nuestra hoja de word
    '''
    documento = Document()

    style = documento.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    section = documento.sections[0]
    section.page_height = Mm(210)
    section.page_width = Mm(297)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    section.orientation = WD_ORIENT.LANDSCAPE

    parrafo_exp = documento.add_paragraph()
    run_exp = parrafo_exp.add_run(nombre_expediente)
    run_exp.bold = True
    run_exp.underline = True
    run_exp.font.size = Pt(16)
    parrafo_exp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return documento



def armar_anexos_dos(planillas,separar):
    '''
    Procesamos todos los archivos con la función armar_anexo_dosV2

    :param planilla: list[archivos] lista de archivos a procesar
    :param separar: bool indica si queremos crear un documento diferente por oficina
    '''

    for planilla in planillas:
        documentos = armar_anexo_dosV2(planilla, separar)
    
    return documentos


##########################################################
##### NO SE USAN, QUEDAN PARA CONSULTAR EN EL FUTURO #####
##########################################################
def validar_archivo_mensualizados(archivo):

    df = pd.read_excel(archivo)
    cant_columnas = len(df.columns)
    valores_nulos = df.iloc[:, [0, 1, 2, 3, 4, 5, 7, 8]].isnull().any().any()

    archivo.seek(0)
    return df.columns, cant_columnas, valores_nulos

def validar_otro_archivo(archivo):

    df = pd.read_excel(archivo)
    cant_columnas = len(df.columns)
    valores_nulos = df.iloc[:, [0,1,2,3,4,5]].isnull().any().any() 
    
    archivo.seek(0)
    return df.columns, cant_columnas, valores_nulos