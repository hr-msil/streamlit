
import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pdfminer.high_level import extract_text
import re
from io import BytesIO



def obtener_expdte(filename):
    expdte_pdf = filename.split(' - ')[1]
    expdte = expdte_pdf.split('.')[0]
    return expdte

def obtener_reso(filename):
    reso = filename.split(' - ')[0]
    return reso

def escribir_acto_admin(documento, reso, exp):

    documento.add_paragraph()
    parrafo_acta = documento.add_paragraph()
    run_acta = parrafo_acta.add_run(f'Se deja en constancia que en el día de la fecha los agentes fueron notificados de la resolución {reso} contenida en el expediente {exp} mediante la cual se _______________ a los mismos.')
    parrafo_acta.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    parrafo_firma = documento.add_paragraph()
    run_firma = parrafo_firma.add_run("FIRMA")
    parrafo_firma.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    parrafo_aclaracion = documento.add_paragraph()
    run_aclaracion = parrafo_aclaracion.add_run("ACLARACIÓN")
    parrafo_aclaracion = WD_PARAGRAPH_ALIGNMENT.LEFT


def extraer_reso_y_expediente(texto):
    texto = texto.strip()

    # Buscar dónde empieza el nroDeExpediente
    match = re.search(r'\bEX-\d{4}-', texto)
    if not match:
        return None, None

    inicio_expte = match.start()

    # Cortamos sin limpiar a ciegas
    nro_reso = texto[:inicio_expte].rstrip(" -")  # solo eliminamos guiones o espacios al final del reso
    nro_expte = texto[inicio_expte:].lstrip(" -") # solo eliminamos guiones o espacios al principio del expte
    nro_expte = nro_expte.split('.pdf')[0]

    #return nro_reso, nro_expte
    return nro_expte,nro_reso

def obtener_nombres_y_legajos(file):
    trabajadores = {} #key = legajo, clave = nombre completo
    texto = extract_text(file)
    texto = texto.replace('\n', ' ') 
    texto = re.sub(r'\s+', ' ', texto)
    patron = r"(?:(?:Dr\.|Dra\.|Lic\.|Ing\.|Sr\.|Sra\.|Prof\.|Mg\.)\s*)?" \
         r"(?P<nombre>" \
             r"(?:[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+|" \
                 r"[A-ZÁÉÍÓÚÑ]+(?:[-'´’][A-ZÁÉÍÓÚÑ]+)*)" \
             r"(?:(?:,\s*|\s+)" \
                 r"(?:(?:de|del|de\slos|de\slas|da|do|dos|das|la|las|los|y|e)\s+)?" \
                 r"(?:[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+|" \
                     r"[A-ZÁÉÍÓÚÑ]+(?:[-'´’][A-ZÁÉÍÓÚÑ]+)*))" \
             r"{0,5}" \
         r")\s*,?\s*" \
         r"\((?:Legajo|Leg\.?)\s+" \
             r"(?:(?:N°|Nº|No|N\.°|N\.º|Nᵃ|Num\.?)\s*)?" \
             r"(?P<legajo>\d{1,3}(?:\.\d{3})*|\d+)" \
         r"\)"

    coincidencias = re.findall(patron, texto)

    for nombre, legajo in coincidencias:
        legajo = legajo.replace('.', '')
        trabajadores[legajo] = nombre
    
    return trabajadores

def obtener_datos(st_archivos):
    datos_expdtes = {}

    for file in st_archivos:
        st.write(f"Vamos con {file.name}")
        #reso = obtener_reso(file.name)
        #expdte = obtener_expdte(file.name)
        reso,expdte = extraer_reso_y_expediente(file.name)
        trabajadores = obtener_nombres_y_legajos(file)
        datos_expdtes[file.name] = {
                                        "expdte": expdte,
                                        "reso": reso,
                                        "trabajadores": trabajadores
                                    }
    return datos_expdtes
# ACTA NOTIFICACION

def armar_hoja(documento,expediente,resolucion,trabajadores):
     # Encabezado del expediente centrado y en negrita
    parrafo_acta = documento.add_paragraph()
    run_acta = parrafo_acta.add_run("ACTA DE NOTIFICACIÓN")
    run_acta.bold = True
    run_acta.underline = True
    run_acta.font.size = Pt(16)
    parrafo_acta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    parrafo_exp = documento.add_paragraph()
    run_exp = parrafo_exp.add_run(expediente)
    run_exp.bold = True
    run_exp.underline = True
    run_exp.font.size = Pt(16)
    parrafo_exp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Encabezado de la resolución centrado, subrayado y en negrita
    parrafo_res = documento.add_paragraph()
    run_res = parrafo_res.add_run(resolucion)
    run_res.bold = True
    run_res.underline = True
    run_res.font.size = Pt(16)
    parrafo_res.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    

    # Tabla
    tabla = documento.add_table(rows=1, cols=3)
    tabla.style = 'Table Grid'  # Bordes visibles

    encabezado = tabla.rows[0].cells
    encabezado[0].text = "LEGAJO"
    encabezado[1].text = "NOMBRE"
    encabezado[2].text = "FIRMA Y FECHA"

    for legajo, nombre in trabajadores.items():
        fila = tabla.add_row().cells
        fila[0].text = str(legajo)
        fila[1].text = nombre
        fila[2].text = ""  # Vacío para firma y fecha

    escribir_acto_admin(documento, expediente, resolucion)

    documento.add_page_break()

def armar_documento(dict,st_files):
    documento = Document()

    style = documento.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    for file in st_files:
        expdte = dict[file.name]["expdte"]
        reso = dict[file.name]["reso"]
        trabajadores = dict[file.name]["trabajadores"]
        armar_hoja(documento,expdte,reso,trabajadores)
    
    return documento







